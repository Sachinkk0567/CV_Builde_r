import streamlit as st
from google import genai
from config import get_api_key

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

st.set_page_config(
    page_title="AI Resume Builder",
    page_icon="📄",
    layout="wide"
)

st.title("📄 AI Resume Builder")

st.sidebar.title("Settings")

api_key = get_api_key()

if not api_key:
    st.sidebar.warning("⚠️ Gemini API Key not detected!")
    user_key = st.sidebar.text_input(
        "Enter your Gemini API Key",
        type="password",
        help="Get a free key from Google AI Studio (aistudio.google.com)"
    )
    if user_key:
        st.session_state.user_api_key = user_key
        api_key = user_key
        st.rerun()

model_name = st.sidebar.selectbox(
    "Select Model",
    [
        "gemini-flash-latest",
        "gemini-2.0-flash",
        "gemini-pro-latest",
        "gemini-1.5-flash"
    ]
)

resume_type = st.sidebar.selectbox(
    "Resume Type",
    [
        "Fresher",
        "Experienced",
        "ATS Friendly"
    ]
)

photo = st.file_uploader(
    "Upload Profile Photo (Optional)",
    type=["jpg", "jpeg", "png"]
)

if not api_key:
    st.info("💡 Please enter your Gemini API Key in the sidebar or set `GEMINI_API_KEY` in Streamlit Secrets to generate resumes.")

client = None
if api_key:
    try:
        client = genai.Client(api_key=api_key)
    except Exception as e:
        st.sidebar.error(f"Client init error: {e}")

with st.form("resume"):

    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone")
    address = st.text_area("Address")

    objective = st.text_area("Career Objective")

    education = st.text_area("Education")

    skills = st.text_area("Skills")

    projects = st.text_area("Projects")

    experience = st.text_area("Experience")

    certifications = st.text_area("Certifications")

    achievements = st.text_area("Achievements")

    languages = st.text_input("Languages")

    generate = st.form_submit_button("Generate Resume")

if generate:
    if not client:
        st.error("⚠️ Gemini API Key is required to generate resume.")
    else:
        prompt = f"""
Create a professional {resume_type} resume.

Name:
{name}

Email:
{email}

Phone:
{phone}

Address:
{address}

Career Objective:
{objective}

Education:
{education}

Skills:
{skills}

Projects:
{projects}

Experience:
{experience}

Certifications:
{certifications}

Achievements:
{achievements}

Languages:
{languages}

Use proper headings.
Use bullet points.
Return only the resume text.
"""

        with st.spinner("Generating Resume..."):
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt
                )
                resume = response.text or ""

                st.success("Resume Generated Successfully")

                if photo:
                    st.image(photo, width=150)

                st.markdown(resume)

                # Download TXT
                st.download_button(
                    "Download TXT",
                    resume,
                    "resume.txt",
                    "text/plain"
                )

                # Create Word File
                doc = Document()
                doc.add_heading("Resume", level=1)
                doc.add_paragraph(resume)
                doc.save("resume.docx")

                with open("resume.docx", "rb") as file:
                    st.download_button(
                        "Download Word",
                        file,
                        "resume.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                # Create PDF
                styles = getSampleStyleSheet()
                pdf = SimpleDocTemplate("resume.pdf")
                story = []

                for line in resume.split("\n"):
                    story.append(Paragraph(line, styles["BodyText"]))

                pdf.build(story)

                with open("resume.pdf", "rb") as file:
                    st.download_button(
                        "Download PDF",
                        file,
                        "resume.pdf",
                        "application/pdf"
                    )

            except Exception as err:
                err_msg = str(err)
                if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg:
                    st.error("⚠️ Gemini API quota limit reached. Please wait a moment or try another model.")
                elif "404" in err_msg or "NOT_FOUND" in err_msg:
                    st.error(f"⚠️ Model '{model_name}' was not found. Please pick another model from the sidebar.")
                else:
                    st.error(f"⚠️ Error generating resume: {err_msg}")